from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colors ──────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0C, 0x0E, 0x1A)
INDIGO    = RGBColor(0x63, 0x66, 0xF1)
LAVENDER  = RGBColor(0xA7, 0x8B, 0xFA)
CYAN      = RGBColor(0x67, 0xE8, 0xF9)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
OFFWHITE  = RGBColor(0xED, 0xE9, 0xFE)
MUTED     = RGBColor(0x94, 0x8F, 0xB8)
DARK_CARD = RGBColor(0x16, 0x18, 0x2E)
AMBER     = RGBColor(0xFB, 0xBF, 0x24)
GREEN     = RGBColor(0x6E, 0xE7, 0xB7)

def set_cell_bg(cell, color_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_para_shading(para, color_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    pPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=11, color=None, font='Outfit'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = color
    return run

def section_rule(doc, color_hex='6366F1'):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(12)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_heading(doc, text, level=1):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(20 if level == 1 else 14)
    para.paragraph_format.space_after = Pt(4)
    if level == 1:
        add_run(para, text, bold=True, size=22, color=LAVENDER)
        section_rule(doc, 'A78BFA')
    elif level == 2:
        add_run(para, text, bold=True, size=15, color=INDIGO)
    elif level == 3:
        add_run(para, text, bold=True, size=12, color=CYAN)

def add_body(doc, text, color=OFFWHITE, size=11):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(8)
    add_run(para, text, size=size, color=color)
    return para

def add_bullet(doc, text, indent=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(4)
    add_run(para, text, size=10.5, color=OFFWHITE)

def add_highlight_box(doc, title, body, title_color=LAVENDER, bg='1A1C35'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6.5)

    # Title
    tp = cell.add_paragraph()
    tp.paragraph_format.space_before = Pt(4)
    tp.paragraph_format.space_after = Pt(2)
    tp.paragraph_format.left_indent = Inches(0.15)
    add_run(tp, title, bold=True, size=11, color=title_color)

    # Body
    bp = cell.add_paragraph()
    bp.paragraph_format.space_before = Pt(0)
    bp.paragraph_format.space_after = Pt(6)
    bp.paragraph_format.left_indent = Inches(0.15)
    add_run(bp, body, size=10, color=OFFWHITE)

    # Remove default first paragraph
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

def make_feature_table(doc, rows_data):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    set_cell_bg(hdr[0], '6366F1')
    set_cell_bg(hdr[1], '6366F1')
    for cell, text in zip(hdr, ['Feature', 'Description']):
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.1)
        add_run(p, text, bold=True, size=10, color=WHITE)

    for i, (feat, desc) in enumerate(rows_data):
        row = table.add_row().cells
        bg = '16182E' if i % 2 == 0 else '0C0E1A'
        set_cell_bg(row[0], bg)
        set_cell_bg(row[1], bg)
        p0 = row[0].paragraphs[0]
        p0.paragraph_format.left_indent = Inches(0.1)
        add_run(p0, feat, bold=True, size=10, color=LAVENDER)
        p1 = row[1].paragraphs[0]
        p1.paragraph_format.left_indent = Inches(0.1)
        add_run(p1, desc, size=10, color=OFFWHITE)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ── Build document ─────────────────────────────────────────────────────────────
doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.85)
section.bottom_margin = Inches(0.85)

# Background — DOCX doesn't support true page bg, use a cover table
cover = doc.add_table(rows=1, cols=1)
cover.alignment = WD_TABLE_ALIGNMENT.CENTER
c = cover.cell(0, 0)
set_cell_bg(c, '0C0E1A')
c.width = Inches(6.5)

title_p = c.paragraphs[0]
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(28)
title_p.paragraph_format.space_after = Pt(4)
add_run(title_p, '❤  CareCompanion', bold=True, size=34, color=LAVENDER)

sub_p = c.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after = Pt(6)
add_run(sub_p, 'AI-Powered Cancer Care Companion', size=14, color=CYAN)

tag_p = c.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag_p.paragraph_format.space_before = Pt(0)
tag_p.paragraph_format.space_after = Pt(28)
add_run(tag_p, 'For patients, caregivers, and families navigating cancer treatment', italic=True, size=11, color=MUTED)

doc.add_paragraph().paragraph_format.space_after = Pt(20)


# ── 1. The Problem ─────────────────────────────────────────────────────────────
add_heading(doc, '1. The Problem')
add_body(doc, 'Cancer care is overwhelming — not just medically, but logistically. Patients and their families are managing dozens of medications across multiple doctors, interpreting lab results, fighting insurance denials, tracking prior authorizations, and coordinating care across entire healthcare systems.', color=OFFWHITE)
add_body(doc, 'Most of this critical information lives scattered across paper records, patient portal logins, voicemails, and memory. There is no single place that holds it all, no assistant that remembers it, and no guide to navigate the system.', color=OFFWHITE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
add_run(p, 'CareCompanion is built to solve exactly this.', bold=True, size=11, color=LAVENDER)


# ── 2. What Is CareCompanion ───────────────────────────────────────────────────
add_heading(doc, '2. What Is CareCompanion')
add_body(doc, 'CareCompanion is an AI-powered mobile and web application designed specifically for cancer patients and their caregivers. It centralizes every aspect of the care journey — medications, appointments, lab results, insurance, symptoms — into a single intelligent companion that remembers everything, proactively alerts you to what needs attention, and helps you navigate the healthcare system with confidence.', color=OFFWHITE)

add_highlight_box(doc,
    '❤  Mission',
    'To make cancer care less overwhelming — giving patients and families the information, guidance, and emotional support they need, exactly when they need it.',
    title_color=LAVENDER
)

add_body(doc, 'CareCompanion is not a generic health app. It is built ground-up for the unique complexity of oncology care — cycles of chemotherapy, rotating specialists, dense lab panels, insurance appeals, and the emotional weight of a diagnosis.', color=OFFWHITE)


# ── 3. Who It's For ───────────────────────────────────────────────────────────
add_heading(doc, '3. Who It\'s For')

add_heading(doc, 'Cancer Patients', level=2)
add_body(doc, 'Individuals living with a cancer diagnosis who want to take ownership of their care, understand their treatment, and never miss a medication or appointment.', color=OFFWHITE)

add_heading(doc, 'Family Caregivers', level=2)
add_body(doc, 'Spouses, children, siblings, and friends who manage care for a loved one — often coordinating multiple providers, navigating insurance, and making high-stakes decisions without medical training.', color=OFFWHITE)

add_heading(doc, 'Care Teams', level=2)
add_body(doc, 'Families can invite multiple members with role-based access so everyone stays informed — whether you\'re the primary caregiver or checking in from across the country.', color=OFFWHITE)


# ── 4. Core Features ──────────────────────────────────────────────────────────
add_heading(doc, '4. Core Features')

add_heading(doc, 'AI Companion Chat', level=2)
add_body(doc, 'The heart of CareCompanion is an intelligent chat interface powered by Claude AI. Users can ask anything — about their medications, upcoming appointments, lab results, insurance coverage, or how they\'re feeling — and get accurate, empathetic, personalized responses.', color=OFFWHITE)
add_body(doc, 'The AI is not a generic chatbot. It knows your patient\'s name, cancer type, stage, treatment phase, current medications, recent labs, and care history. Every response is grounded in your actual situation.', color=OFFWHITE)

add_heading(doc, 'Multi-Agent Intelligence', level=3)
add_body(doc, 'Behind the chat is a multi-agent AI system with six specialist agents that activate based on the nature of each question:', color=OFFWHITE)
specialists = [
    ('Medication Specialist', 'Drug interactions, dosing schedules, refill tracking, side effect guidance'),
    ('Insurance Navigator', 'Claims status, denial appeals, prior authorizations, cost estimation'),
    ('Scheduling Coordinator', 'Appointment prep, post-visit notes, follow-up tracking'),
    ('Wellness Monitor', 'Symptom tracking, caregiver emotional support, daily check-ins'),
    ('Lab Analyst', 'Result interpretation, trend analysis, abnormal value flagging'),
    ('General Companion', 'Profile management, document analysis, health summaries'),
]
make_feature_table(doc, specialists)

add_heading(doc, 'Long-Term Memory', level=2)
add_body(doc, 'CareCompanion remembers everything across every conversation. It extracts key facts from each interaction — medications added, allergies mentioned, preferences expressed, concerns raised — and stores them permanently. The AI references this memory to give responses that feel like talking to someone who truly knows your situation.', color=OFFWHITE)

add_highlight_box(doc,
    '📋  What It Remembers',
    'Medications and dosages • Allergies and reactions • Diagnoses and conditions • Insurance details • Care team preferences • Family roles • Emotional patterns • Prior visit notes',
    title_color=CYAN, bg='0F1120'
)

add_heading(doc, 'Medication Management', level=2)
add_body(doc, 'Track every medication with dose, frequency, prescribing doctor, pharmacy contact, and refill date. The AI monitors for upcoming refills and alerts you before you run out. It can also check for drug interactions across your full medication list.', color=OFFWHITE)

add_heading(doc, 'Lab Results', level=2)
add_body(doc, 'Upload or connect lab results and get plain-English explanations of what each value means, whether it\'s trending in the right direction, and what questions to ask your oncologist. No medical degree required.', color=OFFWHITE)

add_heading(doc, 'Insurance Navigation', level=2)
add_body(doc, 'Insurance is one of the most stressful parts of cancer care. CareCompanion tracks claims, flags denials, drafts appeal letters, monitors prior authorization expiry dates, and estimates out-of-pocket costs — all from a simple chat interface.', color=OFFWHITE)

add_heading(doc, 'Proactive Alerts', level=2)
add_body(doc, 'CareCompanion doesn\'t wait for you to ask. It proactively monitors your care situation and sends notifications for:', color=OFFWHITE)
alerts = [
    'Medication refills due within 3 days',
    'Appointments the next day (with prep prompts)',
    'Prior authorizations expiring soon',
    'New abnormal lab results',
    'Low FSA/HSA balances',
    'Upcoming care milestones',
]
for a in alerts:
    add_bullet(doc, a)

add_heading(doc, 'Symptom Journal', level=2)
add_body(doc, 'Daily check-ins capture how the patient is feeling — pain level, mood, sleep quality, energy, and appetite. Over time, this builds a symptom history that can be shared with the care team and analyzed by the AI to spot patterns or flag concerns.', color=OFFWHITE)

add_heading(doc, 'Doctor Visit Prep & Summary', level=2)
add_body(doc, 'Before every appointment, CareCompanion generates a personalized prep sheet: current medications, recent labs, questions to ask, things to bring. After the visit, capture notes, medication changes, referrals, and follow-ups — all saved to the patient\'s record.', color=OFFWHITE)

add_heading(doc, 'Health Summary Export', level=2)
add_body(doc, 'One tap generates a comprehensive health summary — everything a new doctor, specialist, or emergency room needs to know about the patient. Designed to be printed or shared instantly.', color=OFFWHITE)

add_heading(doc, 'Emergency Info Card', level=2)
add_body(doc, 'A single screen designed for paramedics and ER nurses. Shows the patient\'s name, age, allergies (highlighted prominently), current medications, conditions, insurance, and emergency contact. Accessible even when the patient cannot communicate.', color=OFFWHITE)

add_heading(doc, 'Care Team Sharing', level=2)
add_body(doc, 'Multiple family members can be added to a patient\'s care circle with role-based access — Owner, Editor, or Viewer. Everyone stays on the same page with a shared activity feed.', color=OFFWHITE)

add_heading(doc, 'HealthKit Integration', level=2)
add_body(doc, 'On iOS, CareCompanion connects to Apple Health to automatically sync relevant health data — steps, heart rate, sleep, and more — providing a fuller picture of how the patient is doing day to day.', color=OFFWHITE)


# ── 5. The Mobile Experience ──────────────────────────────────────────────────
add_heading(doc, '5. The Mobile Experience')
add_body(doc, 'CareCompanion has a native iOS app built alongside the web platform. The mobile experience is designed for the real moments of caregiving — in the waiting room, at the pharmacy, during a late-night worry.', color=OFFWHITE)

features_mobile = [
    ('Home Dashboard', 'At-a-glance view of medications, next appointment, wellness, and AI CTA'),
    ('AI Chat', 'Full conversational AI with voice input support'),
    ('Care Hub', 'Family command center with care radar, medication status, and insights'),
    ('Care Timeline', 'Chronological view of all care events — medications, appointments, milestones'),
    ('Scan', 'Camera-based document scanning to extract data from prescriptions, labs, and insurance cards'),
    ('Emergency Card', 'One-tap access to critical patient info for first responders'),
]
make_feature_table(doc, features_mobile)

add_body(doc, 'The design is premium and intentional — dark indigo backgrounds, glowing purple accents, smooth animations, haptic feedback — because people managing cancer care deserve beautiful, calming tools, not clinical interfaces that add to the stress.', color=OFFWHITE)


# ── 6. Privacy & Trust ────────────────────────────────────────────────────────
add_heading(doc, '6. Privacy & Trust')
add_body(doc, 'Health data is among the most sensitive information a person can share. CareCompanion is built with this responsibility at its core.', color=OFFWHITE)

trust_items = [
    ('End-to-End Encryption', 'All health data encrypted in transit and at rest'),
    ('HIPAA-Aware Architecture', 'Designed with HIPAA compliance principles from the ground up'),
    ('Your Data, Your Control', 'You own your data. It is never sold or shared with third parties.'),
    ('Role-Based Access', 'Granular control over who in your family can see and edit what'),
    ('Audit Logging', 'All PHI access is logged for compliance and transparency'),
    ('Session Security', 'Secure token-based authentication with CSRF protection'),
]
make_feature_table(doc, trust_items)


# ── 7. The Vision ─────────────────────────────────────────────────────────────
add_heading(doc, '7. The Vision')
add_body(doc, 'Cancer touches nearly 20 million people in the United States alone. Behind each diagnosis is a patient, and behind each patient is a family doing their best to hold everything together.', color=OFFWHITE)

add_highlight_box(doc,
    '✨  The Opportunity',
    'CareCompanion exists at the intersection of AI capability and human need. The technology to truly help people manage complex, high-stakes health situations now exists — and it should be in the hands of the people who need it most.',
    title_color=LAVENDER
)

add_body(doc, 'The long-term vision is a platform that grows with the patient across their entire care journey — from diagnosis through treatment, remission, and beyond. A companion that learns your specific situation over time, anticipates your needs, advocates for you in the healthcare system, and ensures that no critical detail ever falls through the cracks.', color=OFFWHITE)
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(8)
add_run(p2, 'CareCompanion is not just an app. It is the care system that everyone deserves but most people never get.', bold=True, size=11, color=LAVENDER)


# ── Footer ─────────────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(footer_p, 'carecompanionai.org  •  Built with ❤ for the cancer community', size=9, color=MUTED, italic=True)

# ── Save ───────────────────────────────────────────────────────────────────────
out = '/Users/aryanmotgi/Desktop/CareCompanion_Concept.docx'
doc.save(out)
print(f'Saved: {out}')
